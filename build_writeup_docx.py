"""Convert WRITEUP.md to WRITEUP.docx using python-docx.

Targets the specific markdown shape of WRITEUP.md (headings 1-4, bullets,
tables with alignment, bold/italic/inline-code spans, code blocks,
links rendered as text + URL).
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm

ROOT = Path(__file__).parent
SRC = ROOT / "WRITEUP.md"
DST = ROOT / "WRITEUP.docx"

# ---------- inline span parsing ----------

INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_runs(paragraph, text: str, *, base_bold: bool = False) -> None:
    """Add runs to paragraph, honoring **bold**, *italic*, `code`, [text](url)."""
    if not text:
        return
    # Strip trailing whitespace but keep internal spacing
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        elif part.startswith("["):
            m = LINK_RE.match(part)
            if m:
                label, url = m.group(1), m.group(2)
                r = paragraph.add_run(label)
                r.font.color.rgb = RGBColor(0x1A, 0x5F, 0xB4)
                r.underline = True
            else:
                paragraph.add_run(part)
        else:
            r = paragraph.add_run(part)
            if base_bold:
                r.bold = True


# ---------- table parsing ----------

def is_table_separator(line: str) -> bool:
    """Match | --- | :---: | etc."""
    s = line.strip()
    if not s.startswith("|"):
        return False
    cells = [c.strip() for c in s.strip("|").split("|")]
    return all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= n_cols:
                continue
            cell = table.rows[i].cells[j]
            # Clear default paragraph
            p = cell.paragraphs[0]
            p.text = ""
            add_runs(p, cell_text, base_bold=(i == 0))
            for r in p.runs:
                r.font.size = Pt(9.5)
    # Spacing after table
    doc.add_paragraph()


# ---------- main converter ----------

def convert():
    src_text = SRC.read_text(encoding="utf-8")
    lines = src_text.split("\n")

    doc = Document()

    # Page setup — match Google Docs default-ish margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Base style tightening
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    i = 0
    while i < len(lines):
        line = lines[i]

        # --- horizontal rule ---
        if line.strip() == "---":
            p = doc.add_paragraph()
            p_pr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "888888")
            pbdr.append(bottom)
            p_pr.append(pbdr)
            i += 1
            continue

        # --- headings ---
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            h = doc.add_heading(level=level)
            add_runs(h, text)
            # tighten heading colour
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x0E, 0x2E, 0x5C)
            i += 1
            continue

        # --- code fence ---
        if line.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.4)
            r = p.add_run("\n".join(code_lines))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            continue

        # --- table ---
        if line.strip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            header = split_row(line)
            i += 2  # skip header + separator
            body = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                body.append(split_row(lines[i]))
                i += 1
            add_table(doc, [header] + body)
            continue

        # --- bullet ---
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:].strip())
            i += 1
            continue

        # --- numbered list ---
        if re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_runs(p, text)
            i += 1
            continue

        # --- blockquote ---
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            add_runs(p, line[2:].strip())
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # --- blank line ---
        if not line.strip():
            i += 1
            continue

        # --- paragraph (collect until blank / structural) ---
        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if (
                not nxt.strip()
                or nxt.startswith("#")
                or nxt.startswith("```")
                or nxt.startswith("- ")
                or nxt.startswith("* ")
                or nxt.startswith("> ")
                or nxt.strip() == "---"
                or nxt.strip().startswith("|")
                or re.match(r"^\d+\.\s+", nxt)
            ):
                break
            para_lines.append(nxt)
            i += 1
        para_text = " ".join(l.strip() for l in para_lines)
        p = doc.add_paragraph()
        add_runs(p, para_text)

    doc.save(DST)
    print(f"Wrote {DST} ({DST.stat().st_size:,} bytes)")


if __name__ == "__main__":
    convert()
