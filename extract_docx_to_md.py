"""Extract WRITEUP.docx → WRITEUP_FROM_DOCX.md, preserving:
- Headings (h1-h6) → # markers
- Bold / italic / monospace runs → **/_/`
- Tables → | ... | rows with separator
- Bullet & numbered lists → - / 1.
- Hyperlinks → [text](url) (best-effort; falls back to text if URL missing)
- Block-level order (paragraphs, tables, lists)

This is a one-way sync. After user reviews, WRITEUP_FROM_DOCX.md becomes
the new WRITEUP.md.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

ROOT = Path(__file__).parent
SRC = ROOT / "WRITEUP.docx"
DST = ROOT / "WRITEUP_FROM_DOCX.md"


# ---------- helpers ----------

def _run_text_with_format(run) -> str:
    """Return run text wrapped in markdown emphasis markers."""
    text = run.text
    if not text:
        return ""
    # detect monospace via font name
    is_mono = bool(run.font and run.font.name and
                   any(m in (run.font.name or "").lower()
                       for m in ("consolas", "courier", "mono", "menlo")))
    bold = bool(run.bold)
    italic = bool(run.italic)
    if is_mono:
        return f"`{text}`"
    out = text
    if bold and italic:
        out = f"***{out}***"
    elif bold:
        out = f"**{out}**"
    elif italic:
        out = f"*{out}*"
    return out


def _hyperlinks_in_paragraph(paragraph: Paragraph, doc: Document) -> dict[str, str]:
    """Map run-id-ish to URL. python-docx doesn't expose hyperlinks directly;
    walk the XML."""
    out: dict[str, str] = {}
    p_el = paragraph._p
    for hl in p_el.iter(qn("w:hyperlink")):
        rId = hl.get(qn("r:id"))
        if rId and rId in doc.part.rels:
            url = doc.part.rels[rId].target_ref
            for t in hl.iter(qn("w:t")):
                text = t.text or ""
                if text:
                    out[text] = url
    return out


def _paragraph_to_md(paragraph: Paragraph, doc: Document) -> str:
    """Convert a single Paragraph to a markdown line."""
    style = paragraph.style.name if paragraph.style else ""
    text_runs: list[str] = []
    links = _hyperlinks_in_paragraph(paragraph, doc)

    for run in paragraph.runs:
        rt = run.text or ""
        if rt in links:
            # wrap as link
            text_runs.append(f"[{rt}]({links[rt]})")
        else:
            text_runs.append(_run_text_with_format(run))
    body = "".join(text_runs).rstrip()

    # heading detection
    m = re.match(r"^Heading (\d+)$", style)
    if m:
        level = int(m.group(1))
        return f"{'#' * level} {body}"
    if style == "Title":
        return f"# {body}"

    # list detection — python-docx exposes pPr/numPr
    p_pr = paragraph._p.find(qn("w:pPr"))
    is_list = False
    is_numbered = False
    if p_pr is not None:
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is not None:
            is_list = True
            # numId 1 usually = bullet, others may be numbered; not perfectly
            # reliable. Fall back to checking style name.
            if "Number" in style:
                is_numbered = True
    if "List Bullet" in style:
        is_list, is_numbered = True, False
    if "List Number" in style:
        is_list, is_numbered = True, True

    if is_list:
        if is_numbered:
            return f"1. {body}"
        return f"- {body}"

    return body


def _table_to_md(table) -> list[str]:
    """Convert a docx table to a markdown table (lossy on cell merges)."""
    rows: list[list[str]] = []
    for tr in table.rows:
        row_cells = []
        for cell in tr.cells:
            # cell may have multiple paragraphs; join them with <br>
            parts = []
            for p in cell.paragraphs:
                txt = "".join(_run_text_with_format(r) for r in p.runs).strip()
                if txt:
                    parts.append(txt)
            row_cells.append(" ".join(parts) or "")
        rows.append(row_cells)
    if not rows:
        return []
    n_cols = max(len(r) for r in rows)
    out: list[str] = []
    header = rows[0] + [""] * (n_cols - len(rows[0]))
    out.append("| " + " | ".join(header) + " |")
    out.append("|" + "---|" * n_cols)
    for r in rows[1:]:
        r = r + [""] * (n_cols - len(r))
        out.append("| " + " | ".join(r) + " |")
    out.append("")  # blank line after
    return out


# ---------- main walk ----------

def extract():
    doc = Document(str(SRC))
    body = doc.element.body

    # We walk the body in document order, treating tables and paragraphs
    # uniformly so the markdown reflects the visual order.
    out_lines: list[str] = []
    p_iter = iter(doc.paragraphs)
    t_iter = iter(doc.tables)

    # Build a position map: element → (kind, obj)
    children_in_order = []
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            try:
                children_in_order.append(("p", next(p_iter)))
            except StopIteration:
                pass
        elif tag == "tbl":
            try:
                children_in_order.append(("t", next(t_iter)))
            except StopIteration:
                pass

    prev_was_list = False
    for kind, obj in children_in_order:
        if kind == "p":
            line = _paragraph_to_md(obj, doc)
            # collapse a hard \n between same-type list lines
            out_lines.append(line)
            prev_was_list = line.lstrip().startswith(("- ", "1. "))
        elif kind == "t":
            out_lines.append("")  # ensure blank line before table
            out_lines.extend(_table_to_md(obj))

    # collapse 3+ blank lines into 2
    text = "\n".join(out_lines)
    text = re.sub(r"\n{3,}", "\n\n", text).strip() + "\n"

    DST.write_text(text, encoding="utf-8")
    print(f"Wrote {DST} ({len(text):,} chars, {len(text.split()):,} words)")


if __name__ == "__main__":
    extract()
